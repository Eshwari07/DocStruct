from __future__ import annotations

import datetime as _dt
from pathlib import Path
from typing import List, Optional

from docx import Document as DocxDocument  # type: ignore[import]

from docstruct.core.schema import (
    DocumentTree,
    DocNode,
    PageRange,
    SourceFormat,
    ExtractionPath,
)
from docstruct.parsers.base import BaseParser


HEADING_STYLE_TO_DEPTH = {
    "heading 1": 1,
    "heading 2": 2,
    "heading 3": 3,
    "heading 4": 4,
    "heading 5": 5,
    "heading 6": 6,
    "heading 7": 7,
    "heading 8": 8,
    "heading 9": 9,
    "title": 1,
    "subtitle": 2,
}


def _paragraph_depth(p) -> Optional[int]:
    style = getattr(p, "style", None)
    name = getattr(style, "name", None)
    if not name:
        return None
    return HEADING_STYLE_TO_DEPTH.get(name.lower())


def _approximate_pages(doc: DocxDocument) -> dict[int, int]:
    """
    Very lightweight page approximation using XML page break markers.
    Returns mapping paragraph_index -> physical_page (1-based).
    """
    pages: dict[int, int] = {}
    page_num = 1
    for idx, para in enumerate(doc.paragraphs):
        pages[idx] = page_num
        # Inspect underlying XML for page breaks
        elm = para._p
        if elm.xpath(".//w:pageBreak | .//w:lastRenderedPageBreak"):
            page_num += 1
    return pages


class DocxParser(BaseParser):
    """
    Fast-path DOCX parser using paragraph styles for headings.
    """

    def parse(self) -> DocumentTree:
        doc = DocxDocument(self.file_path)
        page_map = _approximate_pages(doc)
        total_pages = max(page_map.values()) if page_map else 1

        root_nodes: List[DocNode] = []
        stack: List[tuple[int, DocNode]] = []

        current_node: Optional[DocNode] = None

        for idx, para in enumerate(doc.paragraphs):
            depth = _paragraph_depth(para)
            text = para.text.strip()

            if depth is not None and text:
                # Heading paragraph
                page = page_map.get(idx, 1)
                pages = PageRange(
                    physical_start=page,
                    physical_end=page,
                    logical_start=None,
                    logical_end=None,
                )

                node = DocNode(
                    title=text,
                    text="",
                    pages=pages,
                    confidence=1.0,
                    images=[],
                    children=[],
                )

                while stack and stack[-1][0] >= depth:
                    stack.pop()

                if stack:
                    stack[-1][1].add_child(node)
                else:
                    root_nodes.append(node)

                stack.append((depth, node))
                current_node = node
            else:
                # Body paragraph, attach to current heading
                if current_node is None:
                    continue
                if text:
                    if current_node.text:
                        current_node.text += "\n\n" + text
                    else:
                        current_node.text = text

        tree = DocumentTree(
            source_file=str(self.file_path.name),
            source_format=SourceFormat.DOCX,
            total_pages=total_pages,
            extracted_at=_dt.datetime.utcnow().isoformat() + "Z",
            extraction_path=ExtractionPath.FAST,
            nodes=root_nodes,
        )
        tree.finalise()
        return tree

